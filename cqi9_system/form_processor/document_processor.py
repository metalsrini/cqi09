"""
Document Processor
=================

This module provides functionality for processing different types of documents,
including PDF, Word documents, and Excel spreadsheets.
"""

import logging
import os
from typing import Dict, List, Any, Optional, Union, Tuple, BinaryIO
from enum import Enum
import json

import pypdf
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import pandas as pd
import docx
from openpyxl import load_workbook

from ..config.config import active_config

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    """Enum representing supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    CSV = "csv"
    JSON = "json"
    UNKNOWN = "unknown"


class DocumentProcessor:
    """
    Processes different types of documents to extract structured information.
    
    This class handles the parsing and extraction of data from various document formats
    commonly used in CQI-9 assessments, including PDFs, Word documents, and spreadsheets.
    """
    
    def __init__(self, ocr_enabled: bool = True):
        """
        Initialize the Document Processor.
        
        Args:
            ocr_enabled: Whether to use OCR for text extraction from images in PDFs.
        """
        self.ocr_enabled = ocr_enabled
        
    @staticmethod
    def detect_document_type(file_path: str) -> DocumentType:
        """
        Detect the type of document based on file extension.
        
        Args:
            file_path: Path to the document file.
            
        Returns:
            The detected document type.
        """
        _, ext = os.path.splitext(file_path)
        ext = ext.lower().lstrip('.')
        
        if ext == 'pdf':
            return DocumentType.PDF
        elif ext == 'docx':
            return DocumentType.DOCX
        elif ext == 'xlsx':
            return DocumentType.XLSX
        elif ext == 'csv':
            return DocumentType.CSV
        elif ext == 'json':
            return DocumentType.JSON
        else:
            return DocumentType.UNKNOWN
            
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document file and extract its content.
        
        Args:
            file_path: Path to the document file.
            
        Returns:
            Dictionary containing the extracted document content and metadata.
        """
        try:
            doc_type = self.detect_document_type(file_path)
            
            if doc_type == DocumentType.PDF:
                return self.process_pdf(file_path)
            elif doc_type == DocumentType.DOCX:
                return self.process_docx(file_path)
            elif doc_type == DocumentType.XLSX:
                return self.process_xlsx(file_path)
            elif doc_type == DocumentType.CSV:
                return self.process_csv(file_path)
            elif doc_type == DocumentType.JSON:
                return self.process_json(file_path)
            else:
                logger.warning(f"Unsupported document type for file: {file_path}")
                return {
                    "success": False,
                    "error": "Unsupported document type",
                    "file_path": file_path
                }
                
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "file_path": file_path
            }
            
    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Process a PDF document and extract its content.
        
        Args:
            file_path: Path to the PDF file.
            
        Returns:
            Dictionary containing the extracted PDF content and metadata.
        """
        try:
            # Extract text using PyPDF
            pdf_content = {"pages": []}
            
            with open(file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                
                # Extract document info
                pdf_content["metadata"] = {
                    "title": pdf_reader.metadata.title if pdf_reader.metadata else None,
                    "author": pdf_reader.metadata.author if pdf_reader.metadata else None,
                    "subject": pdf_reader.metadata.subject if pdf_reader.metadata else None,
                    "creator": pdf_reader.metadata.creator if pdf_reader.metadata else None,
                    "producer": pdf_reader.metadata.producer if pdf_reader.metadata else None,
                    "page_count": len(pdf_reader.pages)
                }
                
                # Extract text from each page
                for i, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    
                    # If text extraction failed or returned very little text and OCR is enabled,
                    # try using OCR
                    if self.ocr_enabled and (not text or len(text.strip()) < 100):
                        logger.info(f"Using OCR for page {i+1} of {file_path}")
                        try:
                            # Convert PDF page to image
                            images = convert_from_path(file_path, first_page=i+1, last_page=i+1)
                            if images:
                                # Extract text using OCR
                                text = pytesseract.image_to_string(images[0])
                        except Exception as ocr_error:
                            logger.error(f"OCR failed for page {i+1} of {file_path}: {str(ocr_error)}")
                    
                    # Add page content
                    pdf_content["pages"].append({
                        "page_number": i + 1,
                        "text": text or ""
                    })
            
            pdf_content["success"] = True
            pdf_content["file_path"] = file_path
            return pdf_content
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "file_path": file_path
            }
            
    def process_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Process a Word document and extract its content.
        
        Args:
            file_path: Path to the Word document file.
            
        Returns:
            Dictionary containing the extracted Word document content and metadata.
        """
        try:
            doc = docx.Document(file_path)
            
            # Extract document properties
            doc_properties = {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
            # Extract paragraphs
            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():  # Only include non-empty paragraphs
                    paragraphs.append({
                        "index": i,
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal"
                    })
            
            # Extract tables
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                    
                tables.append({
                    "index": i,
                    "rows": len(table.rows),
                    "columns": len(table.rows[0].cells) if table.rows else 0,
                    "data": table_data
                })
            
            return {
                "success": True,
                "file_path": file_path,
                "metadata": doc_properties,
                "content": {
                    "paragraphs": paragraphs,
                    "tables": tables
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "file_path": file_path
            }
            
    def process_xlsx(self, file_path: str) -> Dict[str, Any]:
        """
        Process an Excel spreadsheet and extract its content.
        
        Args:
            file_path: Path to the Excel file.
            
        Returns:
            Dictionary containing the extracted Excel content and metadata.
        """
        try:
            # Load workbook
            workbook = load_workbook(filename=file_path, data_only=True)
            
            # Get workbook properties
            wb_properties = {
                "sheet_names": workbook.sheetnames,
                "sheet_count": len(workbook.sheetnames)
            }
            
            # Process each sheet
            sheets = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Convert sheet to pandas DataFrame
                data = []
                for row in sheet.iter_rows(values_only=True):
                    data.append(row)
                    
                # Create DataFrame if there's data
                if data:
                    df = pd.DataFrame(data[1:], columns=data[0] if data else None)
                    
                    # Convert to dictionary
                    sheet_data = {
                        "name": sheet_name,
                        "rows": sheet.max_row,
                        "columns": sheet.max_column,
                        "headers": list(df.columns) if not df.empty else [],
                        "data": df.fillna("").to_dict(orient="records") if not df.empty else []
                    }
                else:
                    sheet_data = {
                        "name": sheet_name,
                        "rows": 0,
                        "columns": 0,
                        "headers": [],
                        "data": []
                    }
                    
                sheets.append(sheet_data)
            
            return {
                "success": True,
                "file_path": file_path,
                "metadata": wb_properties,
                "sheets": sheets
            }
            
        except Exception as e:
            logger.error(f"Error processing Excel file {file_path}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "file_path": file_path
            }
            
    def process_csv(self, file_path: str) -> Dict[str, Any]:
        """
        Process a CSV file and extract its content.
        
        Args:
            file_path: Path to the CSV file.
            
        Returns:
            Dictionary containing the extracted CSV content and metadata.
        """
        try:
            # Read CSV file
            df = pd.read_csv(file_path)
            
            # Get basic metadata
            metadata = {
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": list(df.columns)
            }
            
            # Convert to dictionary
            data = df.fillna("").to_dict(orient="records")
            
            return {
                "success": True,
                "file_path": file_path,
                "metadata": metadata,
                "data": data
            }
            
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "file_path": file_path
            }
            
    def process_json(self, file_path: str) -> Dict[str, Any]:
        """
        Process a JSON file and extract its content.
        
        Args:
            file_path: Path to the JSON file.
            
        Returns:
            Dictionary containing the extracted JSON content.
        """
        try:
            # Read JSON file
            with open(file_path, 'r') as f:
                data = json.load(f)
            
            # Get basic metadata
            if isinstance(data, list):
                metadata = {
                    "type": "array",
                    "count": len(data)
                }
            elif isinstance(data, dict):
                metadata = {
                    "type": "object",
                    "keys": list(data.keys())
                }
            else:
                metadata = {
                    "type": type(data).__name__
                }
            
            return {
                "success": True,
                "file_path": file_path,
                "metadata": metadata,
                "data": data
            }
            
        except Exception as e:
            logger.error(f"Error processing JSON file {file_path}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "file_path": file_path
            } 